import os
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

WORKDIR = "/home/wakkii/PythonProjects/csv_view"

# -----------------------------------------------------------------------------
# 1. Graphviz Dot File Generation
# -----------------------------------------------------------------------------
DOT_CONTENT = """digraph AppFlowchart {
    graph [
        label="📊 CSVデータビジュアルアナライザー & 高度解析ダッシュボード (app.py) 処理フローチャート"
        labelloc="t"
        labeljust="c"
        fontsize=18
        fontname="Noto Sans CJK JP"
        bgcolor="#FFFFFF"
        rankdir=TB
        nodesep=0.4
        ranksep=0.5
        splines=ortho
    ];

    node [
        fontname="Noto Sans CJK JP"
        fontsize=10
        shape=rect
        style="filled,rounded"
        fillcolor="#F8F9FA"
        color="#D3D3D3"
        penwidth=1.2
        margin="0.2,0.1"
    ];

    edge [
        fontname="Noto Sans CJK JP"
        fontsize=9
        color="#595959"
        penwidth=1.2
        arrowsize=0.8
    ];

    // --- Start & Init ---
    start [label="🚀 アプリケーション起動\\n(Streamlit Server)", shape=ellipse, fillcolor="#1F4E79", fontcolor="#FFFFFF", penwidth=0];
    
    subgraph cluster_init {
        label = "1. 初期化・ページ設定";
        style = "dashed";
        color = "#2E75B6";
        fontcolor = "#2E75B6";
        fontname = "Noto Sans CJK JP";
        fontsize = 11;
        
        page_config [label="st.set_page_config()\\n・タイトル/アイコン設定\\n・レイアウト: wide\\n・サイドバー: expanded", fillcolor="#EBF3FA"];
        css_inject [label="st.markdown(<style>...)\\n・カスタムCSS注入\\n(.stMetric, .insight-card, .info-card)", fillcolor="#EBF3FA"];
    }

    // --- Data Source ---
    subgraph cluster_input {
        label = "2. データソース選択 & 読み込み (load_csv_data)";
        style = "dashed";
        color = "#4B6B94";
        fontcolor = "#4B6B94";
        fontname = "Noto Sans CJK JP";
        fontsize = 11;

        mode_select [label="1. 作業モードの選択\\n(図のみ / 分析のみ / 両方)", shape=box, fillcolor="#F0F4F8"];
        source_select [label="2. データソースの選択\\n(サンプルCSV / ユーザーアップロード)", shape=box, fillcolor="#F0F4F8"];
        cond_source [label="データソース種類?", shape=diamond, fillcolor="#FFF2CC", color="#D6B656"];
        
        load_sample [label="サンプルデータ読み込み\\n(filepath='sample_data.csv')", fillcolor="#E2EFDA"];
        load_upload [label="st.file_uploader()\\nユーザー指定CSVファイル取得", fillcolor="#E2EFDA"];
        
        load_func [label="【関数】 load_csv_data()\\n順次エンコーディング試行\\n(UTF-8 -> Shift-JIS -> CP932 -> EUC-JP)\\nPandas DataFrame生成", fillcolor="#D9E1F2", shape=component];
    }

    cond_df [label="DataFrame正常取得?\\n(df is not None and not df.empty)", shape=diamond, fillcolor="#FCE4D6", color="#C55A11"];
    stop_app [label="🛑 案内メッセージ表示 & st.stop()\\n(処理一時停止)", shape=ellipse, fillcolor="#F8CECC", fontcolor="#B85450"];

    // --- Data Profiling ---
    subgraph cluster_profiling {
        label = "3. データプロファイリング (自動列分類)";
        style = "dashed";
        color = "#70AD47";
        fontcolor = "#70AD47";
        fontname = "Noto Sans CJK JP";
        fontsize = 11;

        col_classify [label="列データ型の識別\\n・全項目列 (all_columns)\\n・数値列 (numeric_columns)\\n・カテゴリ列 (categorical_columns)\\n・日付列 (date_columns: 60%以上日付変換可)", fillcolor="#E2EFDA"];
    }

    // --- Sidebar Params ---
    subgraph cluster_sidebar_params {
        label = "4. グラフ・分析パラメータ設定 (サイドバー)";
        style = "dashed";
        color = "#ED7D31";
        fontcolor = "#ED7D31";
        fontname = "Noto Sans CJK JP";
        fontsize = 11;

        param_select [label="軸・グループの選択\\n・X軸列 (x_column)\\n・Y軸数値列 (y_columns: 複数選択可)\\n・グループ分け/色分け列 (group_column)", fillcolor="#FCE4D6"];
        chart_type_select [label="表示グラフ種別の選択\\n(折れ線/棒/散布/ヒストグラム/箱ひげ/円)", fillcolor="#FCE4D6"];
    }

    // --- Main Header ---
    main_header [label="メイン画面表示\\n・ヘッダー & サマリーメトリクス (行数/列数/数値列/日付列)\\n・データプレビュータブ (データ一覧/基本統計/欠損値)", fillcolor="#EDEDED"];

    // --- Tab Branching ---
    tab_branch [label="作業モード分岐\\n(app_mode)", shape=diamond, fillcolor="#FFF2CC", color="#D6B656"];

    // --- Modules ---
    subgraph cluster_mod1 {
        label = "モジュール 1: 📊 インタラクティブ可視化";
        style = "filled";
        fillcolor="#F2F4F8";
        color="#4F46E5";
        fontcolor="#4F46E5";
        fontname="Noto Sans CJK JP";
        fontsize=11;

        mod1_proc [label="Plotly可視化描画\\n・選択グラフを2列レイアウトで生成\\n・plotly_white テーマ適用\\n・折れ線/棒/散布/ヒスト/箱ひげ/円", fillcolor="#E0E7FF"];
    }

    subgraph cluster_mod2 {
        label = "モジュール 2: 🧪 統計学的分析";
        style = "filled";
        fillcolor="#F0FDF4";
        color="#16A34A";
        fontcolor="#16A34A";
        fontname="Noto Sans CJK JP";
        fontsize=11;

        mod2_sub1 [label="2-1. 詳細記述統計量\\n平均/中央値/標準偏差/IQR/歪度/尖度/95%信頼区間", fillcolor="#DCFCE7"];
        mod2_sub2 [label="2-2. 相関分析\\nピアソン/スピアマン相関行列・ヒートマップ・p値有意性", fillcolor="#DCFCE7"];
        mod2_sub3 [label="2-3. 仮説検定・グループ比較\\n2群: Welchのt検定 / Mann-Whitney U\\n多群: ANOVA / Kruskal-Wallis + 箱ひげ図", fillcolor="#DCFCE7"];
    }

    subgraph cluster_mod3 {
        label = "モジュール 3: 📐 線形分析 (回帰分析)";
        style = "filled";
        fillcolor="#EFF6FF";
        color="#2563EB";
        fontcolor="#2563EB";
        fontname="Noto Sans CJK JP";
        fontsize=11;

        mod3_proc [label="OLS最小二乗法モデル (statsmodels)\\n・目的変数 Y vs 説明変数 X\\n・R² / 調整済みR² / RMSE\\n・回帰係数・p値テーブル & 回帰方程式\\n・単回帰直線 / 予測値 vs 実測値 & 残差プロット", fillcolor="#DBEAFE"];
    }

    subgraph cluster_mod4 {
        label = "モジュール 4: ⏳ 時系列分析";
        style = "filled";
        fillcolor="#FEF3C7";
        color="#D97706";
        fontcolor="#D97706";
        fontname="Noto Sans CJK JP";
        fontsize=11;

        mod4_proc [label="時系列解析・トレンド推計\\n・datetime変換 & 昇順ソート\\n・移動平均 (SMA) & ボリンジャーバンド (±2σ)\\n・前期比成長率 (%) & 累積和\\n・線形トレンドモデルによる将来予測 (LinearRegression)", fillcolor="#FDE68A"];
    }

    footer [label="💡 フッター解説 & 利用ガイド表示", fillcolor="#EDEDED"];
    end_node [label="🏁 処理完了 (ユーザー操作待機)", shape=ellipse, fillcolor="#1F4E79", fontcolor="#FFFFFF", penwidth=0];

    // --- Connections ---
    start -> page_config;
    page_config -> css_inject;
    css_inject -> mode_select;
    mode_select -> source_select;
    source_select -> cond_source;

    cond_source -> load_sample [label="サンプルデータ"];
    cond_source -> load_upload [label="自分のCSV"];

    load_sample -> load_func;
    load_upload -> load_func;

    load_func -> cond_df;
    cond_df -> stop_app [label="No (空/無効)"];
    cond_df -> col_classify [label="Yes (正常)"];

    col_classify -> param_select;
    param_select -> chart_type_select;
    chart_type_select -> main_header;

    main_header -> tab_branch;

    tab_branch -> mod1_proc [label="図を表示"];
    tab_branch -> mod2_sub1 [label="データを分析"];
    tab_branch -> mod3_proc [label="データを分析"];
    tab_branch -> mod4_proc [label="データを分析"];
    tab_branch -> mod1_proc [label="両方を表示"];
    tab_branch -> mod2_sub1 [label="両方を表示"];
    
    mod2_sub1 -> mod2_sub2 -> mod2_sub3;

    mod1_proc -> footer;
    mod2_sub3 -> footer;
    mod3_proc -> footer;
    mod4_proc -> footer;

    footer -> end_node;
}
"""

def generate_graphviz():
    dot_path = os.path.join(WORKDIR, "app_flowchart.dot")
    png_path = os.path.join(WORKDIR, "app_flowchart.png")
    
    with open(dot_path, "w", encoding="utf-8") as f:
        f.write(DOT_CONTENT)
    
    cmd = f"dot -Tpng -Gdpi=200 {dot_path} -o {png_path}"
    res = subprocess.run(cmd, shell=True, capture_output=True, text=True)
    if res.returncode == 0:
        print(f"Successfully generated flowchart image at: {png_path}")
    else:
        print(f"Error compiling graphviz: {res.stderr}")

# -----------------------------------------------------------------------------
# 2. Word (.docx) Document Generation
# -----------------------------------------------------------------------------
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_body_paragraph(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Meiryo"
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F5F7FA")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="18" w:space="0" w:color="5C6B73"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_cell_text(cell, text, bold=False, font_size=9.5, color=RGBColor(0x33, 0x33, 0x33), align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Meiryo"
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    return run

def generate_docx():
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("app.py プログラム処理フローチャート解説書")
    run_title.font.name = "Meiryo"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(16)
    run_sub = p_sub.add_run("CSVデータビジュアルアナライザー & 高度解析ダッシュボード 全体処理フロー構造定義書")
    run_sub.font.name = "Meiryo"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    run_div = p_div.add_run("―" * 55)
    run_div.font.color.rgb = RGBColor(0xD3, 0xD3, 0xD3)

    # 1. 概要
    add_heading_1(doc, "1. はじめに")
    add_body_paragraph(doc, "本ドキュメントは、`app.py`（StreamlitベースのCSV可視化・統計解析Webアプリケーション）の全コード処理フローを図解および詳細な説明表としてまとめた設計仕様ドキュメントです。")
    add_body_paragraph(doc, "アプリケーションの初期化、データ読み込み・文字コード自動判別、サイドバー制御、データプロファイリング、および各種可視化・分析モジュールへの分岐処理の流れを体系化しています。")

    # 2. フローチャート図解
    add_heading_1(doc, "2. 全体フローチャート図")
    add_body_paragraph(doc, "以下は `app.py` の実行フローを可視化したビジュアルフローチャートです。")

    png_path = os.path.join(WORKDIR, "app_flowchart.png")
    if os.path.exists(png_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(6)
        p_img.add_run().add_picture(png_path, width=Inches(6.2))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(12)
        r_cap = p_cap.add_run("▲ 図 1.1: app.py 全体処理フローチャート")
        r_cap.font.name = "Meiryo"
        r_cap.font.size = Pt(9.0)
        r_cap.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # 3. 処理ステップ詳細仕様表
    add_heading_1(doc, "3. 処理ステップ詳細仕様一覧")
    add_body_paragraph(doc, "フローチャート内の各処理ブロックの役割、該当行番号、および主な処理内容は以下の通りです。")

    table = doc.add_table(rows=10, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B0C4DE", sz="6")
    
    col_widths = [Inches(1.1), Inches(1.8), Inches(1.5), Inches(2.1)]
    headers = ["フェーズ", "処理ブロック名", "該当行番号", "詳細仕様・主な関数/ロジック"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=150, right=150)
        format_cell_text(hdr_cells[i], h, bold=True, font_size=10, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

    steps_data = [
        ("1. 初期化", "ページ構成 & CSS設定", "40行目～80行目", "st.set_page_config() による画面レイアウト指定、st.markdown() でカスタムCSS注入"),
        ("2. データ読み込み", "CSVローダー関数", "86行目～120行目", "load_csv_data(): UTF-8, Shift-JIS, CP932, EUC-JP を順次試行しDataFrame化"),
        ("3. UI制御", "サイドバー設定入力", "125行目～167行目", "作業モード選択（図/分析/両方）およびデータソース選択（サンプル/アップロード）"),
        ("4. 判定・検証", "データ有無判定ガード", "169行目～172行目", "df が empty または None の場合、案内メッセージを出力して st.stop() で処理中断"),
        ("5. 解析準備", "データ自動プロファイリング", "176行目～193行目", "全列、数値列、カテゴリ列、および日付判定列（60%以上認識可能）の動的識別"),
        ("6. パラメータ設定", "軸・グラフ種別選択", "198行目～243行目", "X軸、Y軸（複数選択）、グループ分け列の指定、各種グラフの表示ON/OFF設定"),
        ("7. 画面表示", "ヘッダー & プレビュー", "248行目～278行目", "メトリクス表示（行数/列数等）、データ一覧・記述統計・欠損値チェックのタブ表示"),
        ("8. タブ分岐", "作業モード動的タブ切り替え", "283行目～300行目", "選択された作業モード（app_mode）に基づき表示するタブコンポーネントを構築"),
        ("9. モジュール実行", "4大分析モジュール処理", "307行目～822行目", "【モジュール1】図表示（Plotly）\n【モジュール2】統計分析（記述統計/相関/検定）\n【モジュール3】線形分析（OLS回帰）\n【モジュール4】時系列（移動平均/ボリンジャー/予測）")
    ]

    for row_idx, data in enumerate(steps_data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell_value in enumerate(data):
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
            bold = True if col_idx in [0, 1] else False
            format_cell_text(row_cells[col_idx], cell_value, bold=bold, font_size=9, align=align)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. Mermaidダイアグラムコード定義
    add_heading_1(doc, "4. Mermaid フローチャート構文定義")
    add_body_paragraph(doc, "ドキュメントツールやGitHub markdown等で再利用可能な Mermaid 形式のフローチャートコードです。")

    mermaid_code = """graph TD
    A[🚀 アプリ起動 Streamlit] --> B[1. ページ基本設定 st.set_page_config]
    B --> C[2. カスタムCSS注入 st.markdown]
    C --> D[3. サイドバー: 作業モード & データソース選択]
    D --> E{データソース選択?}
    E -- サンプルCSV --> F1[sample_data.csv 読み込み]
    E -- ユーザー指定 --> F2[st.file_uploader]
    F1 --> G[load_csv_data 関数: エンコーディング自動試行]
    F2 --> G
    G --> H{DataFrame 正常取得?}
    H -- No 空またはエラー --> I[st.stop 処理中断]
    H -- Yes 正常読み込み --> J[4. データ自動プロファイリング 列分類]
    J --> K[5. サイドバー: 軸・パラメータ・グラフ選択]
    K --> L[6. メインヘッダー & データプレビュー表示]
    L --> M{作業モード分岐 app_mode}
    M -- 📈 図を表示する --> N1[モジュール1: Plotlyインタラクティブ可視化]
    M -- 🧪 データを分析する --> N2[モジュール2: 統計学的分析]
    M -- 🧪 データを分析する --> N3[モジュール3: OLS線形分析 回帰分析]
    M -- 🧪 データを分析する --> N4[モジュール4: 時系列分析 & 将来予測]
    M -- 📊 両方を表示 --> N1 & N2 & N3 & N4
    N1 --> O[💡 フッター解説表示]
    N2 --> O
    N3 --> O
    N4 --> O
    O --> P[🏁 処理完了 ユーザー操作待機]"""
    
    add_code_block(doc, mermaid_code)

    docx_path = os.path.join(WORKDIR, "app_flowchart.docx")
    doc.save(docx_path)
    print(f"Successfully generated docx file at: {docx_path}")

# -----------------------------------------------------------------------------
# 3. Markdown (.md) Document Generation
# -----------------------------------------------------------------------------
def generate_markdown():
    md_content = """# 📊 app.py プログラム処理フローチャート & 構造解説書

本ドキュメントは、StreamlitベースのCSVデータ解析Webアプリケーション [`app.py`](file:///home/wakkii/PythonProjects/csv_view/app.py) の処理フローをビジュアルダイアグラムおよび詳細なステップ表で定義した設計仕様書です。

---

## 1. ビジュアルフローチャート (全体構成図)

![app.py フローチャート](app_flowchart.png)

---

## 2. Mermaid フローチャートダイアグラム

```mermaid
graph TD
    A[🚀 アプリ起動 Streamlit] --> B[1. ページ基本設定 st.set_page_config]
    B --> C[2. カスタムCSS注入 st.markdown]
    C --> D[3. サイドバー: 作業モード & データソース選択]
    D --> E{データソース選択?}
    E -- サンプルCSV --> F1[sample_data.csv 読み込み]
    E -- ユーザー指定 --> F2[st.file_uploader]
    F1 --> G[load_csv_data 関数: エンコーディング自動試行]
    F2 --> G
    G --> H{DataFrame 正常取得?}
    H -- No 空またはエラー --> I[st.stop 処理中断]
    H -- Yes 正常読み込み --> J[4. データ自動プロファイリング 列分類]
    J --> K[5. サイドバー: 軸・パラメータ・グラフ選択]
    K --> L[6. メインヘッダー & データプレビュー表示]
    L --> M{作業モード分岐 app_mode}
    M -- 📈 図を表示する --> N1[モジュール1: Plotlyインタラクティブ可視化]
    M -- 🧪 データを分析する --> N2[モジュール2: 統計学的分析]
    M -- 🧪 データを分析する --> N3[モジュール3: OLS線形分析 回帰分析]
    M -- 🧪 データを分析する --> N4[モジュール4: 時系列分析 & 将来予測]
    M -- 📊 両方を表示 --> N1 & N2 & N3 & N4
    N1 --> O[💡 フッター解説表示]
    N2 --> O
    N3 --> O
    N4 --> O
    O --> P[🏁 処理完了 ユーザー操作待機]
```

---

## 3. 処理フェーズ別 詳細説明一覧

| ステップ | 処理フェーズ | 該当コード範囲 | 処理内容と使用ライブラリ / ロジック |
| :--- | :--- | :--- | :--- |
| **Step 1** | **初期化 & ページ設定** | L40 - L80 | `st.set_page_config()` でワイド表示設定。<br>`st.markdown()` でカスタムCSSを注入（`.stMetric`, `.insight-card`, `.info-card` 等の装飾）。 |
| **Step 2** | **データ読み込み処理** | L86 - L120 | `load_csv_data()` 関数。<br>日本語環境の文字化け防止のため `UTF-8` → `Shift-JIS` → `CP932` → `EUC-JP` の順で `pd.read_csv()` を自動順次試行。 |
| **Step 3** | **サイドバーUI制御** | L125 - L167 | 作業モード（可視化/分析/両方）およびデータソース（サンプルデータ/ユーザーアップロード）の受付。 |
| **Step 4** | **データ判定ガード** | L169 - L172 | `df is None or df.empty` の場合、エラーメッセージを表示して `st.stop()` で後続処理をストップ。 |
| **Step 5** | **データプロファイリング** | L176 - L193 | 読み込まれたDataFrameの各列を型分類（全列、数値列、カテゴリ列）。<br>60%以上の行が日付変換可能な列を `date_columns` として自動識別。 |
| **Step 6** | **軸 & パラメータ設定** | L198 - L243 | サイドバーでX軸、Y軸（複数可）、グループ分け列を設定。<br>図表示モードの場合は表示するグラフ種類（折れ線/棒/散布/ヒスト/箱ひげ/円）をチェックボックスで受領。 |
| **Step 7** | **メイン画面 & プレビュー** | L248 - L278 | サマリーメトリクス（総レコード数、列数等）を表示。<br>タブ切替で「データプレビュー」「基本記述統計サマリー」「欠損値チェック」を表示。 |
| **Step 8** | **モジュール動的タブ分岐** | L283 - L300 | ユーザーが指定した `app_mode` に応じてメイン画面のタブコンポーネントを動的に生成。 |
| **Step 9-1**| **【モジュール1】可視化** | L307 - L399 | `Plotly` によるインタラクティブグラフ描画（`plotly_white` テーマ）。2列グリッド配置で各種グラフを動的レンダリング。 |
| **Step 9-2**| **【モジュール2】統計分析** | L404 - L573 | `scipy.stats` を活用。<br>・**サブタブ2-1**: 詳細記述統計量（信頼区間、歪度、尖度）<br>・**サブタブ2-2**: ピアソン/スピアマン相関ヒートマップ & p値判定<br>・**サブタブ2-3**: 仮説検定（2群: t検定/Mann-Whitney U、多群: ANOVA/Kruskal-Wallis） |
| **Step 9-3**| **【モジュール3】線形分析** | L578 - L685 | `statsmodels.api` OLS最小二乗法回帰分析。<br>決定係数 $R^2$、RMSE、回帰係数・p値テーブル、自動方程式生成、回帰直線 & 残差プロット描画。 |
| **Step 9-4**| **【モジュール4】時系列分析** | L690 - L822 | `scikit-learn LinearRegression` ＆ `pandas`。<br>・移動平均（SMA）＆ ボリンジャーバンド（$\pm 2\sigma$）<br>・前期比成長率（%）＆ 累積和<br>・線形トレンド推計による将来予測グラフ＆データテーブル出力。 |
| **Step 10** | **フッター表示 & 完了** | L827 - L835 | アプリの利用ガイドフッターを出力し、ユーザーの次のインタラクション待機状態へ遷移。 |

---

## 4. 関連ファイル
- 📜 Pythonメインコード: [`app.py`](file:///home/wakkii/PythonProjects/csv_view/app.py)
- 📄 Word形式解説書: [`app_flowchart.docx`](file:///home/wakkii/PythonProjects/csv_view/app_flowchart.docx)
- 🎨 高解像度フローチャート画像: [`app_flowchart.png`](file:///home/wakkii/PythonProjects/csv_view/app_flowchart.png)
"""
    md_path = os.path.join(WORKDIR, "app_flowchart.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"Successfully generated markdown file at: {md_path}")

if __name__ == "__main__":
    generate_graphviz()
    generate_docx()
    generate_markdown()

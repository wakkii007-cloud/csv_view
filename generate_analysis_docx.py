import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """セル背景色を設定"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    """セル内の余白（パディング）を設定（dxa単位: 20 dxa = 1 pt）"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    """表全体の罫線を設定（洗練された水平線主体）"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def make_callout(doc, text_list, title="💡 ポイント・解釈基準", border_color="1F4E79", bg_color="F0F4F8"):
    """洗練されたコールアウトボックスを作成"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    # 左ボーダーのみ強調
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.name = "Meiryo"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    for item in text_list:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.15
        run = p2.add_run(item)
        run.font.name = "Meiryo"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_text, caption=""):
    """コードブロックを追加"""
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(4)
        p_cap.paragraph_format.space_after = Pt(2)
        r_cap = p_cap.add_run(f"【実装コード】 {caption}")
        r_cap.font.name = "Meiryo"
        r_cap.font.size = Pt(9.0)
        r_cap.font.bold = True
        r_cap.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F8F9FA")
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="18" w:space="0" w:color="5C6B73"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.0)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
        
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(4)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_body_p(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Meiryo"
        r_pre.font.size = Pt(10)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def format_cell(cell, text, bold=False, font_size=9.0, color=RGBColor(0x33, 0x33, 0x33), align=WD_ALIGN_PARAGRAPH.LEFT, font_name="Meiryo"):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    return run

def create_styled_table(doc, headers, rows_data, col_widths=None):
    """スタイリングされたデータテーブルを作成"""
    table = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="CCCCCC", sz="4", val="single")
    
    # ヘッダー行
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        format_cell(hdr_cells[i], title, bold=True, font_size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
    # ヘッダープロパティ（改ページ時にもヘッダーを繰り返し表示）
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    # データ行
    for r_idx, row_items in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F9FBFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_items):
            align = WD_ALIGN_PARAGRAPH.LEFT
            font_name = "Meiryo"
            if c_idx == 0 and len(headers) > 3:
                align = WD_ALIGN_PARAGRAPH.LEFT
            elif "数式" in headers[c_idx] or "コード" in headers[c_idx]:
                font_name = "Consolas"
            format_cell(row_cells[c_idx], str(val), bold=False, font_size=8.8, color=RGBColor(0x33, 0x33, 0x33), align=align, font_name=font_name)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=100, right=100)
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        # 行の途中での改ページ禁止
        r_trPr = table.rows[r_idx + 1]._tr.get_or_add_trPr()
        r_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    # カラム幅の設定
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
                
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(6)
    return table

def generate_analysis_specification_docx(output_filename="app_analysis_logic_specification.docx"):
    doc = docx.Document()
    
    # 余白設定 (1インチ = 25.4mm)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # ヘッダー・フッター
        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.text = "app.py 高度データ解析・統計・回帰・時系列ロジック仕様書"
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_hdr.runs[0].font.name = "Meiryo"
        p_hdr.runs[0].font.size = Pt(8.5)
        p_hdr.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.text = "CSV Data Visual Analyzer & Advanced Analytics Dashboard | Confidential"
        p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ftr.runs[0].font.name = "Meiryo"
        p_ftr.runs[0].font.size = Pt(8.5)
        p_ftr.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # =========================================================================
    # 表紙・ドキュメントタイトルエリア
    # =========================================================================
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("📊 app.py 分析ロジック詳細仕様書")
    r_title.font.name = "Meiryo"
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("記述統計・仮説検定・線形回帰・時系列分析の実装アルゴリズムと数理モデル")
    r_sub.font.name = "Meiryo"
    r_sub.font.size = Pt(12)
    r_sub.font.color.rgb = RGBColor(0x5C, 0x6B, 0x73)

    # メタデータ表
    meta_headers = ["項目", "仕様・詳細内容"]
    meta_rows = [
        ["対象ソースコード", "app.py (Streamlit / Python 3.10+)"],
        ["対象分析モジュール", "モジュール2: 統計学的分析 / モジュール3: 線形分析(回帰) / モジュール4: 時系列分析"],
        ["主要解析ライブラリ", "scipy.stats, statsmodels.api, sklearn.linear_model, pandas, numpy"],
        ["ドキュメント作成日", "2026年8月18日"],
        ["文書ステータス", "確定版（設計・アルゴリズム・実装コード完全リファレンス）"]
    ]
    create_styled_table(doc, meta_headers, meta_rows, col_widths=[2.0, 4.5])

    # =========================================================================
    # 第1章：システム概要と分析アーキテクチャ
    # =========================================================================
    add_heading_1(doc, "第1章 システム概要と分析アーキテクチャ")
    add_body_p(doc, "本アプリケーション（app.py）は、CSVファイルを入力として、直感的なデータ可視化機能に加え、学術的・実務的な意思決定を支援する4つの主要な高度データ解析エンジンを統合したStreamlitダッシュボードです。")
    
    add_heading_2(doc, "1.1 分析機能の全体構成")
    add_body_p(doc, "サイドバーで「🧪 データを分析する (高度解析)」または「📊 両方を表示 (全機能)」を選択した際に、以下の3つの主要タブ（4つの分析ドメイン）が動的に展開されます。")

    arch_headers = ["分析区分", "機能タブ", "主要アルゴリズム / 手法", "採用ライブラリ"]
    arch_rows = [
        ["記述統計", "🧪 統計学的分析\n(サブタブ1)", "平均、中央値、標準偏差、分散、IQR、歪度、尖度、95%信頼区間", "pandas, scipy.stats"],
        ["相関分析", "🧪 統計学的分析\n(サブタブ2)", "ピアソン相関係数、スピアマン順位相関係数、無相関検定(p値)", "pandas, scipy.stats, plotly"],
        ["仮説検定", "🧪 統計学的分析\n(サブタブ3)", "2群比較 (Welch t検定, Mann-Whitney U検定)\n多群比較 (一元配置ANOVA, Kruskal-Wallis検定)", "scipy.stats, plotly"],
        ["線形回帰", "📐 線形分析", "OLS通常最小二乗法、R²、調整済みR²、RMSE、t検定、F検定、残差分析", "statsmodels.api, sklearn"],
        ["時系列分析", "⏳ 時系列分析", "単純移動平均 (SMA)、ボリンジャーバンド (±2σ)、前期比成長率、累積和、線形トレンド推計将来予測", "pandas, sklearn, plotly"]
    ]
    create_styled_table(doc, arch_headers, arch_rows, col_widths=[1.1, 1.4, 2.7, 1.3])

    add_heading_2(doc, "1.2 データ前処理パイプライン")
    add_body_p(doc, "分析を安定かつ正確に実行するため、以下の自動前処理パイプラインが組み込まれています：")
    add_body_p(doc, "UTF-8, Shift-JIS, CP932, EUC-JP を順次自動試行し、日本語CSVの文字化け・読み込み失敗を防止。", bold_prefix="① 文字コード自動判別ローダー: ")
    add_body_p(doc, "df.select_dtypes(include=[np.number]) により数値列を抽出。非数値列はカテゴリ・グループ分け列として自動分類。", bold_prefix="② データ型の動的プロファイリング: ")
    add_body_p(doc, "pd.to_datetime によるパース成功率が60%を超える列を日付・時間列として自動検出し、時系列分析の対象として推薦。", bold_prefix="③ 日付型の自動認識: ")
    add_body_p(doc, "各検定・回帰分析の実行直前に .dropna() を適用し、欠損値による計算エラーを完全に隔離。", bold_prefix="④ 欠損値（NaN）の自動除外: ")

    # =========================================================================
    # 第2章：記述統計量および相関分析ロジック
    # =========================================================================
    add_heading_1(doc, "第2章 記述統計量および相関分析ロジック")
    
    add_heading_2(doc, "2.1 詳細記述統計量の算出仕様")
    add_body_p(doc, "数値列ごとに、一般的な基本統計量（平均・中央値等）に加え、分布の歪み（歪度・尖度）や母平均の区間推定（95%信頼区間）を含む全10指標を即座に計算・一覧化します。")

    stat_headers = ["統計指標名", "数理定義式", "実装コード (app.py)", "分析上の意味と解釈"]
    stat_rows = [
        ["サンプル数 (N)", "N = len(x)", "n = len(series)", "有効なデータ行数（欠損値除外後）。"],
        ["算術平均値 (Mean)", "μ = (1/N) Σ x_i", "mean_val = series.mean()", "データの重心となる中心値。外れ値の影響を受けやすい。"],
        ["中央値 (Median)", "第50パーセンタイル値", "median_val = series.median()", "データを昇順に並べた時の中央値。外れ値に対して頑健（ロバスト）。"],
        ["標本標準偏差 (Std)", "s = √[ (1/(N-1)) Σ(x_i - μ)² ]", "std_val = series.std()", "データのばらつき度合い（不偏標準偏差）。"],
        ["標本分散 (Var)", "s² = (1/(N-1)) Σ(x_i - μ)²", "series.var()", "標準偏差の二乗。ばらつきの絶対量。"],
        ["四分位範囲 (IQR)", "IQR = Q3 (75%) - Q1 (25%)", "series.quantile(0.75) - series.quantile(0.25)", "中央50%のデータ幅。箱ひげ図の箱の長さに対応。"],
        ["歪度 (Skewness)", "g1 = m3 / (m2^(3/2))", "skew_val = stats.skew(series)", "左右の非対称性。0: 左右対称、正: 右に裾が長い、負: 左に裾が長い。"],
        ["尖度 (Kurtosis)", "g2 = (m4 / m2²) - 3", "kurt_val = stats.kurtosis(series)", "分布の尖り具合（Fisher定義: 正規分布=0）。正: 鋭利、負: 平坦。"],
        ["標準誤差 (SEM)", "SE = s / √N", "sem_val = stats.sem(series)", "標本平均そのもののばらつき度合い。"],
        ["95% 信頼区間 (CI)", "[ μ - t*SE,  μ + t*SE ]", "stats.t.interval(0.95, df=n-1, loc=mean_val, scale=sem_val)", "母平均が95%の確率で含まれる区間（自由度 N-1 のt分布に基づく推定）。"]
    ]
    create_styled_table(doc, stat_headers, stat_rows, col_widths=[1.3, 1.7, 1.8, 1.7])

    make_callout(doc, [
        "歪度 (Skewness) の目安: |Skewness| < 0.5 はほぼ対称、0.5〜1.0 は中程度の偏り、1.0以上は強い歪みを示します。",
        "尖度 (Kurtosis) の目安: 0より大きければ正規分布よりピークが高く裾が厚い（外れ値が生じやすい）、0より小さければピークが低く平坦な分布です。",
        "95%信頼区間の自由度: サンプルサイズ n > 1 の場合に stats.t.interval を適用。n=1の場合は [mean, mean] を返却しゼロ除算例外を防御しています。"
    ], title="💡 記述統計の判定・解釈基準")

    code_desc_stats = """# app.py: 詳細記述統計の計算ロジック抜粋
series = df[col].dropna()
n = len(series)
mean_val = series.mean()
std_val = series.std()
median_val = series.median()
iqr_val = series.quantile(0.75) - series.quantile(0.25)
skew_val = stats.skew(series)
kurt_val = stats.kurtosis(series)

# 95%信頼区間（t分布に基づく平均値の推定範囲）の計算
sem_val = stats.sem(series)
ci_range = stats.t.interval(0.95, df=n-1, loc=mean_val, scale=sem_val) if n > 1 else (mean_val, mean_val)"""
    add_code_block(doc, code_desc_stats, "詳細記述統計および95%信頼区間の算出 (app.py: 418-431行)")

    add_heading_2(doc, "2.2 相関分析ロジック（ピアソン vs スピアマン & p値判定）")
    add_body_p(doc, "2つ以上の数値列が存在する場合、変数間の相関行列の算出、Plotlyヒートマップ描画、および全変数ペアにおける相関係数と無相関検定（p値）の算出を実施します。")
    
    add_heading_3(doc, "(1) 相関係数の算出手法")
    add_body_p(doc, "線形な比例関係の強さを測定。正規分布を前提とするパラメトリック指標（式: r = Σ(x-x̄)(y-ȳ) / [√Σ(x-x̄)² √Σ(y-ȳ)²]）。", bold_prefix="• ピアソン積率相関 (Pearson): ")
    add_body_p(doc, "各変数を順位（Rank）に変換した上で相関を測定。非線形な単調増加・減少関係や、外れ値を含むデータに対して頑健なノンパラメトリック指標。", bold_prefix="• スピアマン順位相関 (Spearman): ")

    add_heading_3(doc, "(2) 相関の強さおよび有意性判定基準")
    corr_rule_headers = ["判定項目", "判定条件 (コード条件式)", "分類・表示メッセージ"]
    corr_rule_rows = [
        ["非常に強い相関", "abs(r) >= 0.7", "「非常に強い」相関関係（正の相関または負の相関）"],
        ["強い相関", "0.4 <= abs(r) < 0.7", "「強い」相関関係"],
        ["弱い / 相関なし", "abs(r) < 0.4", "「弱い/なし」相関関係"],
        ["統計的有意差あり", "p_val < 0.05", "「有意 (p < 0.05)」: 偶然ではなく有意な関係"],
        ["有意差なし", "p_val >= 0.05", "「有意差なし (p ≥ 0.05)」: 偶然の範囲を出ない"]
    ]
    create_styled_table(doc, corr_rule_headers, corr_rule_rows, col_widths=[1.5, 2.0, 3.0])

    code_corr = """# app.py: 相関分析とp値（有意性）計算ロジック
corr_matrix = df[numeric_columns].corr(method=method_code)

pairs = []
for i in range(len(numeric_columns)):
    for j in range(i+1, len(numeric_columns)):
        col1, col2 = numeric_columns[i], numeric_columns[j]
        val = corr_matrix.loc[col1, col2]
        valid_df = df[[col1, col2]].dropna()
        
        # scipyによる無相関検定とp値の計算
        if method_code == 'pearson':
            _, p_val = stats.pearsonr(valid_df[col1], valid_df[col2])
        else:
            _, p_val = stats.spearmanr(valid_df[col1], valid_df[col2])
        
        pairs.append({
            "変数 1": col1,
            "変数 2": col2,
            "相関係数": round(val, 3),
            "相関の強さ": "非常に強い" if abs(val)>=0.7 else ("強い" if abs(val)>=0.4 else "弱い/なし"),
            "p値 (有意性)": f"{p_val:.4f}",
            "統計的有意差": "有意 (p < 0.05)" if p_val < 0.05 else "有意差なし (p ≥ 0.05)"
        })"""
    add_code_block(doc, code_corr, "相関係数およびp値判定 (app.py: 466-503行)")

    # =========================================================================
    # 第3章：統計的仮説検定・グループ間比較ロジック
    # =========================================================================
    add_heading_1(doc, "第3章 統計的仮説検定・グループ間比較ロジック")
    add_body_p(doc, "カテゴリ列（グループ分け変数）と数値列（目的比較変数）を選択することで、グループ数（2群か3群以上か）を自動判別し、パラメトリック検定とノンパラメトリック検定の両方を並行実行します。")

    add_heading_2(doc, "3.1 2群比較検定ロジック（t検定 / Mann-Whitney U検定）")
    add_body_p(doc, "性別、A/Bテスト、処置群/対照群など、カテゴリ値が2つの場合に自動実行されます。")

    test2_headers = ["検定手法", "統計モデル / 前提条件", "検定統計量と実装関数", "特徴と採用理由"]
    test2_rows = [
        [
            "Welchのt検定\n(パラメトリック)",
            "正規分布を前提とするが、2群間の等分散性は仮定しない。",
            "stats.ttest_ind(g1, g2, equal_var=False)\nt統計量 & 両側p値",
            "Studentのt検定よりも偽陽性（第1種の過誤）を強力に抑制できる現代の標準手法。"
        ],
        [
            "Mann-Whitney U検定\n(ノンパラメトリック)",
            "正規分布を前提とせず、データ全体の順位（ランク和）に基づく検定。",
            "stats.mannwhitneyu(g1, g2)\nU統計量 & 両側p値",
            "サンプル数が少ない場合や、外れ値・歪んだ分布に対しても安全に適用可能。"
        ]
    ]
    create_styled_table(doc, test2_headers, test2_rows, col_widths=[1.5, 1.8, 1.8, 1.4])

    make_callout(doc, [
        "判定基準（有意水準 α = 0.05）: p値 < 0.05 の場合、「統計的に有意差あり（帰無仮説：2群の平均/分布は等しい を棄却）」と判定し、緑色の成功メッセージを表示します。",
        "p値 ≥ 0.05 の場合、「有意差なし（帰無仮説を棄却できない）」と判定し、青色の案内メッセージを表示します。"
    ], title="💡 2群比較の判定ロジック")

    add_heading_2(doc, "3.2 多群比較検定ロジック（一元配置ANOVA / Kruskal-Wallis検定）")
    add_body_p(doc, "地域（東日本/西日本/九州等）や年代別など、グループ数が3つ以上の場合に自動実行されます。")

    testm_headers = ["検定手法", "統計モデル / 前提条件", "検定統計量と実装関数", "判定と結果の解釈"]
    testm_rows = [
        [
            "一元配置分散分析\n(One-way ANOVA)",
            "各群が正規分布に従うことを前提とし、群間分散と群内分散の比率を評価。",
            "stats.f_oneway(*group_data_list)\nF統計量 & p値",
            "p < 0.05 の場合、「いずれかの群間に有意な平均差が存在する」と判定。"
        ],
        [
            "Kruskal-Wallis検定\n(ノンパラメトリックANOVA)",
            "正規分布を仮定せず、全データの順位和を用いて群間の分布のズレを評価。",
            "stats.kruskal(*group_data_list)\nH統計量 & p値",
            "分布形状が非正規である場合や順序尺度データに対しても信頼できる多群検定。"
        ]
    ]
    create_styled_table(doc, testm_headers, testm_rows, col_widths=[1.5, 1.8, 1.8, 1.4])

    code_tests = """# app.py: 2群および多群の統計的仮説検定ロジック抜粋
groups = df[group_var].dropna().unique()

if len(groups) == 2:
    # 2群比較
    g1_data = df[df[group_var] == groups[0]][target_num_var].dropna()
    g2_data = df[df[group_var] == groups[1]][target_num_var].dropna()
    
    t_stat, t_pval = stats.ttest_ind(g1_data, g2_data, equal_var=False)
    u_stat, u_pval = stats.mannwhitneyu(g1_data, g2_data)

elif len(groups) > 2:
    # 3群以上の多群比較
    group_data_list = [df[df[group_var] == g][target_num_var].dropna() for g in groups]
    
    f_stat, f_pval = stats.f_oneway(*group_data_list)
    kw_stat, kw_pval = stats.kruskal(*group_data_list)"""
    add_code_block(doc, code_tests, "仮説検定・グループ比較ロジック (app.py: 523-555行)")

    # =========================================================================
    # 第4章：線形回帰分析（OLS）・要因解析ロジック
    # =========================================================================
    add_heading_1(doc, "第4章 線形回帰分析（OLS）・要因解析ロジック")
    add_body_p(doc, "最小二乗法（Ordinary Least Squares: OLS）を用い、目的変数（Y）に対して1つ以上の説明変数（X）が与える影響度（回帰係数）とモデルの当てはまり度合いを詳細に推計・診断します。")

    add_heading_2(doc, "4.1 OLSモデル推定アルゴリズム")
    add_body_p(doc, "1. 目的変数 Y と 説明変数群 X を抽出し、.dropna() により欠損値行を除外。")
    add_body_p(doc, "2. statsmodels.api.add_constant(X) により、切片項（定数項 const: 列値がすべて1の列）を説明変数行列に付加。")
    add_body_p(doc, "3. sm.OLS(Y, X_with_const).fit() を実行し、回帰パラメータ行列および各種統計量を推定。")

    add_heading_2(doc, "4.2 モデル適合度評価指標の定義と算出")
    reg_eval_headers = ["指標名", "数理定義式", "実装プロパティ (statsmodels / sklearn)", "解釈と評価基準"]
    reg_eval_rows = [
        ["決定係数 (R²)", "R² = 1 - (SS_res / SS_tot)", "model.rsquared", "モデルが目的変数の変動を説明できている割合 (0〜1)。1に近いほど高精度。"],
        ["調整済み R²", "R̄² = 1 - [ (1-R²)(n-1)/(n-p-1) ]", "model.rsquared_adj", "説明変数の増加による見かけ上のR²上昇を補正した値。変数選択の指標。"],
        ["RMSE (二乗平均平方根誤差)", "RMSE = √[ (1/n) Σ(y_i - ŷ_i)² ]", "np.sqrt(mean_squared_error(Y, fitted))", "予測値と実測値の平均的な乖離の大きさ。目的変数と同じ単位。"],
        ["モデル全体の F検定 p値", "H0: 全ての回帰係数=0", "model.f_pvalue", "p < 0.05 であれば、説明変数の少なくとも1つがYに対して有意に影響。"]
    ]
    create_styled_table(doc, reg_eval_headers, reg_eval_rows, col_widths=[1.5, 1.8, 1.8, 1.4])

    add_heading_2(doc, "4.3 回帰係数パラメータテーブルと有意性検定")
    add_body_p(doc, "各説明変数（および定数項）に対して以下のパラメータを算出し、テーブルとして出力します：")
    
    coeff_headers = ["項目名", "算出方法 / プロパティ", "統計学的解釈"]
    coeff_rows = [
        ["回帰係数 (B)", "model.params", "説明変数が1単位増加したときの目的変数（Y）の平均変化量。"],
        ["標準誤差 (SE)", "model.bse", "係数推定量 B のばらつき度合い。"],
        ["t値 (t-statistic)", "model.tvalues (= B / SE)", "係数が0からどれだけ離れているかの比率。"],
        ["p値 (P > |t|)", "model.pvalues", "帰無仮説（係数=0）の検定p値。p < 0.05 で統計的有意。"],
        ["95% 信頼区間", "model.conf_int()", "真の回帰係数が95%の確率で含まれる推定範囲 [下限, 上限]。"]
    ]
    create_styled_table(doc, coeff_headers, coeff_rows, col_widths=[1.5, 2.0, 3.0])

    add_heading_2(doc, "4.4 回帰方程式の自動生成と残差診断プロット")
    add_body_p(doc, "推定された係数に基づき、人間が読みやすい回帰方程式テキストを自動生成します（例: 売上 = 120.5 + 1.45×(広告費) - 0.32×(価格)）。")
    add_body_p(doc, "さらに、以下の2種類の診断グラフを描画し、モデルの前提条件（線形性、残差の等分散性・独立性）を検証します：")
    add_body_p(doc, "単回帰時は散布図＋OLS回帰直線、重回帰時は「実測値 vs 予測値」プロットに45度対角線（完全予測線）を重ねて適合度を可視化。", bold_prefix="① 適合度プロット: ")
    add_body_p(doc, "予測値 ŷ を横軸、残差 e = y - ŷ を縦軸にプロット。e=0 の基準線周りに残差が均一にランダム散布されているか（不均一分散や非線形パターンの有無）を診断。", bold_prefix="② 残差プロット (Residual Plot): ")

    code_reg = """# app.py: 線形回帰モデル学習と回帰方程式生成ロジック抜粋
clean_df = df[[target_y] + features_x].dropna()
Y = clean_df[target_y]
X = clean_df[features_x]
X_with_const = sm.add_constant(X) # 切片（定数項）を追加

# OLS線形モデルの学習
model = sm.OLS(Y, X_with_const).fit()

# 回帰方程式のテキスト自動生成
eq_terms = [f"{model.params['const']:.3f}"] if 'const' in model.params else []
for feat in features_x:
    val = model.params[feat]
    sign = "+" if val >= 0 else "-"
    eq_terms.append(f"{sign} {abs(val):.3f} × ({feat})")
formula_str = f"**{target_y}** = " + " ".join(eq_terms)"""
    add_code_block(doc, code_reg, "OLSモデル学習と回帰式の文字列構築 (app.py: 593-639行)")

    # =========================================================================
    # 第5章：時系列分析・将来予測ロジック
    # =========================================================================
    add_heading_1(doc, "第5章 時系列分析・将来予測ロジック")
    add_body_p(doc, "時系列データに対し、トレンドの平滑化（移動平均・ボリンジャーバンド）、成長推移（前期比・累積和）、および線形トレンド推計に基づく将来予測を実行します。")

    add_heading_2(doc, "5.1 前処理と時系列ソート")
    add_body_p(doc, "選択された日付列を pd.to_datetime(..., errors='coerce') で日時に変換し、変換失敗（NaT）の行を除外した上で、時系列昇順（古い日付から新しい日付へ）に完全ソート（.sort_values('parsed_date')）します。最低3行以上の有効データが必要です。")

    add_heading_2(doc, "5.2 移動平均 (SMA) とボリンジャーバンド (±2σ)")
    add_body_p(doc, "スライダーで指定されたウィンドウ期間 k に基づき、過去 k 期の単純移動平均と移動標準偏差を算出します。")

    ts_calc_headers = ["時系列指標", "算出定義式", "実装コード (app.py)", "実務上の活用"]
    ts_calc_rows = [
        ["単純移動平均 (SMA)", "SMA_t = (1/k) Σ_{i=0}^{k-1} Y_{t-i}", "ts_df[col].rolling(window=k).mean()", "短期的なノイズを除去し、基調となるトレンドを平滑化表示。"],
        ["移動標準偏差 (Std)", "σ_t = √[ (1/(k-1)) Σ(Y_{t-i} - SMA_t)² ]", "ts_df[col].rolling(window=k).std()", "直近 k 期間における価格・数値の変動性（ボラティリティ）。"],
        ["上限バンド (+2σ)", "Upper = SMA_t + 2 × σ_t", "ts_df['sma'] + (2 * ts_df['std'])", "統計的に約95.4%のデータが収まる理論上の上限境界。"],
        ["下限バンド (-2σ)", "Lower = SMA_t - 2 × σ_t", "ts_df['sma'] - (2 * ts_df['std'])", "統計的に約95.4%のデータが収まる理論上の下限境界。"]
    ]
    create_styled_table(doc, ts_calc_headers, ts_calc_rows, col_widths=[1.5, 1.8, 1.8, 1.4])

    add_heading_2(doc, "5.3 前期比成長率 (%) と累積和 (Cumsum)")
    add_body_p(doc, "1期前の値に対する増減率（ts_df[col].pct_change() * 100）および、期首からの累積合計（ts_df[col].cumsum()）を計算し、棒グラフと折れ線グラフで可視化します。")

    add_heading_2(doc, "5.4 線形トレンド推計による将来予測アルゴリズム")
    add_body_p(doc, "過去の推移から長期的なトレンド直線（回帰直線）を機械学習（LinearRegression）により推定し、未来の任意ステップ先（スライダー指定: 1〜10期先）の数値を自動予測・外挿します。")

    make_callout(doc, [
        "ステップ1 (タイムインデックス化): 過去データを時間順に time_idx = [0, 1, 2, ..., N-1] に変換。",
        "ステップ2 (モデル学習): LinearRegression().fit(ts_df[['time_idx']], ts_df[target_ts_col]) により、傾き a と切片 b を学習 (ŷ = a・t + b)。",
        "ステップ3 (未来インデックス外挿): future_time_idx = [N, N+1, ..., N+H-1] を生成し、lr.predict() で将来予測値を計算。",
        "ステップ4 (未来日付の間隔自動補正): 過去データの日付差分の中央値 (ts_df['parsed_date'].diff().median()) を自動算出し、日次・月次・年次等の任意の間隔を正確に保った未来日付リストを生成。"
    ], title="🔮 将来予測アルゴリズムの4ステップ")

    code_forecast = """# app.py: 線形トレンド将来予測ロジック抜粋
ts_df['time_idx'] = np.arange(len(ts_df))
lr = LinearRegression()
lr.fit(ts_df[['time_idx']], ts_df[target_ts_col])

# 未来のインデックスを生成して予測
last_time_idx = ts_df['time_idx'].iloc[-1]
future_time_idx = np.arange(last_time_idx + 1, last_time_idx + 1 + forecast_periods).reshape(-1, 1)
future_preds = lr.predict(future_time_idx)

# データの日付間隔（日/月/年など）の中央値を推定して未来日付を算出
if len(ts_df) > 1:
    date_diff = ts_df['parsed_date'].diff().median()
else:
    date_diff = pd.Timedelta(days=1)

last_date = ts_df['parsed_date'].iloc[-1]
future_dates = [last_date + (i + 1) * date_diff for i in range(forecast_periods)]

forecast_df = pd.DataFrame({
    "parsed_date": future_dates,
    "predicted_val": future_preds
})"""
    add_code_block(doc, code_forecast, "線形トレンド推計と未来日付補間 (app.py: 779-801行)")

    # =========================================================================
    # 第6章：実装コード完全対照・例外処理・境界値リファレンス
    # =========================================================================
    add_heading_1(doc, "第6章 実装コード完全対照・例外処理・境界値リファレンス")
    add_body_p(doc, "app.py における全分析ロジックのソースコード行番号と、例外防御策（フォールトトレランス）の対応一覧です。")

    ref_headers = ["機能項目", "app.py 行番号", "主要関数・API", "例外防御 / 境界値処理対策"]
    ref_rows = [
        ["詳細記述統計", "414 - 455行", "stats.skew, stats.kurtosis, stats.t.interval", "n <= 1 の場合は信頼区間を [mean, mean] としてゼロ除算を防止。dropna() で欠損値を排除。"],
        ["相関分析・p値", "459 - 508行", "df.corr, stats.pearsonr, stats.spearmanr", "ペアごとに valid_df = df[[col1, col2]].dropna() で共通非欠損値のみ抽出しp値計算。"],
        ["2群仮説検定", "526 - 547行", "stats.ttest_ind (Welch), stats.mannwhitneyu", "equal_var=False を明示指定。各群のdropna()を徹底。判定文（結論）を自動切替。"],
        ["多群仮説検定", "548 - 573行", "stats.f_oneway, stats.kruskal", "リスト内包表記で各群データをアンパック (*group_data_list)。3群以上を自動検知。"],
        ["OLS線形回帰", "578 - 686行", "sm.add_constant, sm.OLS, mean_squared_error", "データ件数 > 説明変数数 + 1 の前提をバリデーション。単回帰と重回帰で作図ロジックを自動分岐。"],
        ["時系列分析", "691 - 822行", "pd.to_datetime, rolling, LinearRegression", "errors='coerce'で不正日付をNaT化。有効データ >= 3件を保証。diff().median() で周期を自動推定。"]
    ]
    create_styled_table(doc, ref_headers, ref_rows, col_widths=[1.2, 1.1, 2.0, 2.2])

    doc.save(output_filename)
    print(f"Successfully generated {output_filename} ({os.path.getsize(output_filename):,} bytes)")

if __name__ == "__main__":
    out_file = "app_analysis_logic_specification.docx"
    if len(sys.argv) > 1:
        out_file = sys.argv[1]
    generate_analysis_specification_docx(out_file)

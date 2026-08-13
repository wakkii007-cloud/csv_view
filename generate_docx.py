import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set shading background color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Set internal cell padding (in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    """Set subtle horizontal borders for a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Primary Navy
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6) # Steel Blue
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_body_paragraph(doc, text, bold_prefix="", space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Meiryo"
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run = p.add_run(text)
    run.font.name = "Meiryo"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F5F7FA") # Light neutral background
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    # Left border only (dark slate)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="18" w:space="0" w:color="5C6B73"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    lines = code_text.strip().split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.1
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def make_callout(doc, text_list, title="💡 ポイント", border_color="1F4E79", bg_color="F0F4F8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.name = "Meiryo"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    for item in text_list:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.15
        run = p2.add_run(item)
        run.font.name = "Meiryo"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def format_cell_text(cell, text, bold=False, font_size=9.5, color=RGBColor(0x33, 0x33, 0x33), align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Meiryo"
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    return run

def create_document():
    doc = docx.Document()
    
    # Page setup - Margins 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header Title Block
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("app.py スタイルシート該当箇所・UIデザイン構成 解説書")
    run_title.font.name = "Meiryo"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Primary Navy

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(16)
    run_sub = p_sub.add_run("Streamlitアプリケーション（app.py）におけるCSS定義・レイアウト設定・スタイリング箇所の完全ガイド")
    run_sub.font.name = "Meiryo"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    run_div = p_div.add_run("―" * 55)
    run_div.font.color.rgb = RGBColor(0xD3, 0xD3, 0xD3)

    # 1. はじめに
    add_heading_1(doc, "1. はじめに (概要)")
    add_body_paragraph(doc, "本ドキュメントは、Python環境で動作する Streamlit ベースのWebアプリケーション `app.py` 内において、画面のデザイン、レイアウト、スタイルシート（CSS）の定義およびスタイリング制御が「ファイルのどの部分（行番号・コードブロック）」に存在するのかを体系的に整理・解説した技術ドキュメントです。")
    add_body_paragraph(doc, "`app.py` では、StreamlitのデフォルトUIを拡張し、視認性と操作性を高めるためにカスタムCSSの注入、コンポーネント装飾、およびグラフやテーブルのカラーテーマ設定が施されています。")

    # 2. スタイルシート該当箇所 サマリーテーブル
    add_heading_1(doc, "2. スタイルシート・UIスタイル該当箇所 一覧")
    add_body_paragraph(doc, "`app.py` 全体（全835行）におけるスタイリング関連のコード位置は以下の通りです。")

    # Table creation
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B0C4DE", sz="6")
    
    col_widths = [Inches(1.0), Inches(1.8), Inches(1.8), Inches(2.2)]
    
    # Header Row
    headers = ["行番号", "スタイリング種別", "該当コード・対象要素", "役割・変更内容"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=140, bottom=140, left=150, right=150)
        format_cell_text(hdr_cells[i], h, bold=True, font_size=10, color=RGBColor(0xFF, 0xFF, 0xFF), align=WD_ALIGN_PARAGRAPH.CENTER)

    # Table data
    rows_data = [
        ("40行目～45行目", "ページ基本レイアウト設定", "st.set_page_config()", "画面全体のワイド配置（wide）、ブラウザタイトル、初期サイドバー状態の規定"),
        ("48行目～80行目", "メインカスタムCSS定義\n(スタイルシート注入)", "st.markdown('<style>...', unsafe_allow_html=True)", "余白（padding）調整、メトリクスカード(.stMetric)、インサイトカード(.insight-card)、情報カード(.info-card)のデザイン定義"),
        ("267行目", "データフレーム装飾", "df.describe().T.style\n.highlight_max()", "基本記述統計テーブルにおいて、各項目の最大値セルを動的にハイライト装飾"),
        ("302行目 他", "Plotlyグラフテーマ設定", "theme_template = 'plotly_white'\npx.*(..., template=...) ", "全6種類の可視化グラフおよび時系列・相関グラフに白基調の洗練されたテーマを共通適用"),
        ("639行目", "HTMLクラス適用", "<div class='info-card'>...</div>", "48～80行目で定義したカスタムCSSクラス (.info-card) を回帰方程式表示枠に適用")
    ]

    for row_idx, data in enumerate(rows_data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell_value in enumerate(data):
            row_cells[col_idx].width = col_widths[col_idx]
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            bold = True if col_idx == 1 else False
            format_cell_text(row_cells[col_idx], cell_value, bold=bold, font_size=9, align=align)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 3. 各該当箇所の詳細解説
    add_heading_1(doc, "3. スタイル設定箇所の詳細解説")

    # 3.1 メインのカスタムCSS定義
    add_heading_2(doc, "3.1 メインのカスタムCSS定義ブロック 【行番号: 48行目 ～ 80行目】")
    add_body_paragraph(doc, "アプリケーション内で最も直接的にスタイルシート（CSS）を記述しているのが、48行目から80行目の `st.markdown()` ブロックです。Streamlitでは `unsafe_allow_html=True` オプションを有効にすることで、HTMLの `<style>` タグを介して任意のカスタムCSSをWebページ内に注入することができます。")

    add_heading_3(doc, "▼ 実際のソースコード（48行目～80行目）")
    code_css = """# UI表示をより洗練させるためのカスタムCSSスタイル定義
st.markdown(\"\"\"
    <style>
    /* メインコンテンツエリアの上部・下部余白を調整 */
    .main .block-container {
        padding-top: 1.5rem;
        padding-bottom: 3rem;
    }
    /* メトリクス表示カードの装飾スタイル */
    .stMetric {
        background-color: #f8f9fa;
        padding: 12px;
        border-radius: 8px;
        border-left: 4px solid #4f46e5;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }
    /* 結果・インサイト用カードのスタイル */
    .insight-card {
        background-color: #f0fdf4;
        border-left: 4px solid #16a34a;
        padding: 14px;
        border-radius: 6px;
        margin-bottom: 1rem;
    }
    /* 説明・情報提示用カードのスタイル */
    .info-card {
        background-color: #eff6ff;
        border-left: 4px solid #2563eb;
        padding: 14px;
        border-radius: 6px;
        margin-bottom: 1rem;
    }
    </style>
\"\"\", unsafe_allow_html=True)"""
    add_code_block(doc, code_css)

    add_heading_3(doc, "▼ 各CSSセレクタの役割とスタイル詳細")
    
    make_callout(doc, [
        "1. `.main .block-container`: 画面上部の無駄な空白（デフォルトの大きな余白）を縮小し（padding-top: 1.5rem）、コンテンツ領域のフィット感を高めています。",
        "2. `.stMetric`: Streamlit標準の `st.metric()` ウィジェット（総レコード数などの数値カード）を背景色 #f8f9fa（ライトグレー）、角丸 8px、左端に 4px のインディゴブルー（#4f46e5）のアクセント線とドロップシャドウで高級感のあるカード風デザインに刷新しています。",
        "3. `.insight-card`: 分析結果やインサイトを表示するためのカスタムクラス。背景色 #f0fdf4（爽やかな薄緑）と左端の緑色アクセント線（#16a34a）でポジティブな結果領域をアピールします。",
        "4. `.info-card`: 説明や数式を提示するためのカスタムクラス。背景色 #eff6ff（ソフトな薄青）と左端の青色アクセント線（#2563eb）で視認性を高めたカード構造を作成します。"
    ], title="📌 CSSクラスの詳細分析", border_color="2563EB", bg_color="F0F4FF")

    # 3.2 ページ基本設定
    add_heading_2(doc, "3.2 ページ基本レイアウト設定 【行番号: 40行目 ～ 45行目】")
    add_body_paragraph(doc, "40行目～45行目の `st.set_page_config()` では、Webアプリ全体の画面レイアウト枠組みを指定しています。CSSを適用する前段の構造決定として重要な役割を果たします。")
    
    code_config = """st.set_page_config(
    page_title="CSVデータアナライザー & 高度解析ダッシュボード",
    page_icon="📊",
    layout="wide",                  # 画面全体を広く使うワイドレイアウト
    initial_sidebar_state="expanded" # 起動時にサイドバーを開いた状態にする
)"""
    add_code_block(doc, code_config)
    add_body_paragraph(doc, "・ `layout=\"wide\"`: デフォルトの「中央寄せ・狭い幅」から「画面横幅いっぱいに使うワイド表示」に変更し、ダッシュボードや複数グラフの並列配置を最適化しています。")

    # 3.3 HTMLタグへのCSSクラス適用
    add_heading_2(doc, "3.3 HTMLタグへのCSSクラス適用 【行番号: 639行目】")
    add_body_paragraph(doc, "48～80行目で定義した `.info-card` スタイルシートは、639行目の回帰分析モジュール内で以下のようにHTML要素の class 属性として呼び出されています。")

    code_html_apply = """# 639行目: 回帰方程式のカード表示
st.markdown(f\"<div class='info-card'><b>推定量（推定された回帰方程式）:</b><br>{formula_str}</div>\", unsafe_allow_html=True)"""
    add_code_block(doc, code_html_apply)
    add_body_paragraph(doc, "このように、Streamlitの `st.markdown(..., unsafe_allow_html=True)` を組み合わせることで、Pythonコード内から定義済みCSSクラスを適用した柔軟なUIコンポーネントを描画しています。")

    # 3.4 Plotlyのグラフテーマ
    add_heading_2(doc, "3.4 Plotly グラフコンポーネントのテーマ・配色スタイリング 【行番号: 302行目 他】")
    add_body_paragraph(doc, "データ可視化ライブラリ Plotly を用いたグラフ群全体のビジュアルを統一するため、302行目でグローバルなテーマ変数 `theme_template` が定義されています。")

    code_plotly = """# 302行目: Plotlyの共通グラフテーマ設定
theme_template = "plotly_white"

# 例: 折れ線グラフへのテーマ適用 (327行目)
fig_line = px.line(
    df, x=x_column, y=y_columns, color=color_col,
    title=f"【折れ線】 X: {x_column} vs Y: {', '.join(y_columns)}",
    template=theme_template, markers=True
)"""
    add_code_block(doc, code_plotly)
    add_body_paragraph(doc, "・ `plotly_white`: 背景をシンプルな純白に統一し、グリッド線を薄いグレーに設定することで、データの折れ線や棒グラフのカラーが最も引き立つモダンスタイルを採用しています。")
    add_body_paragraph(doc, "・ その他、相関ヒートマップ（470行目付近）では `color_continuous_scale=\"RdBu_r\"`（赤〜青のグラデーション）や、時系列ボリンジャーバンド（741行目付近）では `fillcolor='rgba(200,200,200,0.2)'`（半透明のバンド領域）など、高度なビジュアルスタイリングが適用されています。")

    # 3.5 Pandas Styler
    add_heading_2(doc, "3.5 Pandas Styler によるテーブル背景ハイライト 【行番号: 267行目】")
    add_body_paragraph(doc, "267行目では、データプレビューエリア内の基本記述統計サマリーテーブルにおいて、各列の最大値を自動的にハイライト表示する Pandas Styler 機能が使用されています。")

    code_styler = """# 267行目: データフレームの最大値ハイライト
st.dataframe(df.describe().T.style.highlight_max(axis=0), use_container_width=True)"""
    add_code_block(doc, code_styler)

    # 4. まとめ・カスタマイズガイド
    add_heading_1(doc, "4. まとめ・デザインカスタマイズ方法")
    add_body_paragraph(doc, "`app.py` におけるスタイルシート設定は、**「メインCSS（48〜80行目）」**、**「レイアウト設定（40〜45行目）」**、**「Plotlyテーマ（302行目）」** の3本柱で構成されています。")

    make_callout(doc, [
        "・ 色を変更したい場合: 48〜80行目の 16進数カラーコード（例: #4f46e5 → お好みのテーマカラー）を書き換えます。",
        "・ 新しいカードデザインを追加したい場合: 48〜80行目の <style> 内に新しいクラス（例: .warning-card { ... }）を追加し、st.markdown(\"<div class='warning-card'>...</div>\", unsafe_allow_html=True) で呼び出します。",
        "・ グラフテーマを変更したい場合: 302行目の theme_template の値を \"plotly_dark\" や \"ggplot2\" に変更することで一括変更が可能です。"
    ], title="🛠️ カスタマイズのアドバイス", border_color="16A34A", bg_color="F0FDF4")

    # Save document
    output_filename = "app_stylesheet_explanation.docx"
    output_path = os.path.join("/home/wakkii/PythonProjects/csv_view", output_filename)
    doc.save(output_path)
    print(f"Successfully generated docx document at: {output_path}")

if __name__ == "__main__":
    create_document()

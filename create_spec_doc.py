import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set shading background color for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set internal cell padding (in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    """Set subtle borders for a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def make_callout(doc, text_list, title="💡 ポイント・注意事項", border_color="1F4E79", bg_color="F0F4F8"):
    """Create a styled callout box with a left border and light shading."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border only
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.name = "Meiryo"
    run_t.font.size = Pt(10.5)
    run_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    for item in text_list:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.15
        run = p2.add_run(item)
        run.font.name = "Meiryo"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def format_cell_text(cell, text, bold=False, font_size=9.5, color=RGBColor(0x33, 0x33, 0x33), align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Meiryo"
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    return run

def create_document():
    doc = docx.Document()
    
    # Page Setup - Margins 1 inch (25.4mm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Normal Style setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Meiryo'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ---------------------------------------------------------
    # COVER / HEADER TITLE
    # ---------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("CSVデータビジュアルアナライザー &\n高度データ解析ダッシュボード")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = 'Meiryo'
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run("プログラム仕様書 (対象ソースコード: app.py)")
    sub_run.font.size = Pt(14)
    sub_run.font.name = 'Meiryo'
    sub_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # Meta Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("ドキュメント名", "システム機能仕様書・プログラム仕様書"),
        ("対象ファイル", "app.py (Streamlit Webアプリケーション)"),
        ("作成日", "2026年8月12日"),
        ("バージョン", "v1.0.0"),
        ("作成主体", "Google Antigravity AI Assistant")
    ]
    set_table_borders(meta_table, color="B0C4DE", sz="6")
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        set_cell_background(cell_lbl, "F0F4F8")
        set_cell_background(cell_val, "FFFFFF")
        set_cell_margins(cell_lbl, top=80, bottom=80, left=120, right=120)
        set_cell_margins(cell_val, top=80, bottom=80, left=120, right=120)
        format_cell_text(cell_lbl, label, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
        format_cell_text(cell_val, val, bold=False)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(18)

    # Helper for Headings
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Meiryo'
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Meiryo'
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Meiryo'
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        return p

    def add_p(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = 'Meiryo'
        run = p.add_run(text)
        run.font.name = 'Meiryo'
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.bold = True
            r_pre.font.name = 'Meiryo'
        run = p.add_run(text)
        run.font.name = 'Meiryo'
        return p

    # ---------------------------------------------------------
    # 1. システム概要
    # ---------------------------------------------------------
    add_h1("第1章 システム概要")
    
    add_h2("1.1 開発目的・背景")
    add_p("本アプリケーション（app.py）は、ユーザーが手元のCSVデータをプログラミング知識なしでアップロードまたは選択するだけで、データの多角的なインタラクティブ可視化から、統計学的仮説検定、線形回帰分析、時系列予測までを一元的に実行できるWebダッシュボードアプリケーションです。")
    add_p("従来の表計算ソフトや専用統計ソフトにおける操作の複雑さを解消し、Webブラウザ上で直感的にデータ探索・意思決定（EDA: Exploratory Data Analysis）を行える環境を提供することを目的としています。")

    add_h2("1.2 主な機能構成")
    add_bullet("ユーザーが保持する独自のCSVファイルまたは同梱のサンプルデータのロード機能", "① データロード & 文字化け防止機能: ")
    add_bullet("データの総行数・列数、数値列、日付列を自動検出してサマリー表示", "② データプロファイリング機能: ")
    add_bullet("折れ線、棒、散布図、ヒストグラム、箱ひげ図、円グラフの即時生成", "③ インタラクティブデータ可視化: ")
    add_bullet("詳細記述統計（信頼区間・歪度・尖度含む）、相関ヒートマップ、仮説検定（t検定/Mann-Whitney, ANOVA/Kruskal-Wallis）", "④ 統計学的分析: ")
    add_bullet("最小二乗法（OLS）に基づく単回帰・重回帰分析、決定係数・RMSEの算出、自動方程式生成、残差分析", "⑤ 線形分析（回帰分析）: ")
    add_bullet("移動平均（SMA）、ボリンジャーバンド（±2σ）、前期比成長率、累積和、線形トレンド将来予測", "⑥ 時系列分析: ")

    add_h2("1.3 動作環境および技術スタック")
    add_p("本アプリケーションはPython環境上で動作し、以下のオープンソースライブラリにより構成されています。")

    tech_table = doc.add_table(rows=7, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tech_table)
    headers = ["ライブラリ / モジュール", "バージョン要件", "主要用途・役割"]
    hdr_cells = tech_table.rows[0].cells
    for i, title in enumerate(headers):
        set_cell_background(hdr_cells[i], "1F4E79")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        format_cell_text(hdr_cells[i], title, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    tech_data = [
        ("Streamlit", ">= 1.30.0", "Web UIフレームワーク、サイドバー制御、タブ・カードレイアウト"),
        ("Pandas", ">= 2.0.0", "CSVデータの読み込み、データフレーム操作、集計・前処理"),
        ("Plotly (Express / Graph Objects)", ">= 5.18.0", "インタラクティブグラフ描画、ヒートマップ、時系列プロット"),
        ("SciPy (stats)", ">= 1.10.0", "記述統計量（歪度・尖度・信頼区間）、仮説検定（t検定, ANOVA, Mann-Whitney, Kruskal-Wallis）"),
        ("Statsmodels", ">= 0.14.0", "OLS最小二乗法回帰分析、モデル統計量・信頼区間・残差計算"),
        ("Scikit-Learn", ">= 1.3.0", "線形回帰モデル (LinearRegression)、決定係数 (R²)・RMSE算出")
    ]
    for row_idx, data_tuple in enumerate(tech_data, start=1):
        row_cells = tech_table.rows[row_idx].cells
        bg_color = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data_tuple):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            format_cell_text(row_cells[col_idx], text, bold=(col_idx==0))
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ---------------------------------------------------------
    # 2. システム構成・動作仕様
    # ---------------------------------------------------------
    add_h1("第2章 システム構成・動作仕様")

    add_h2("2.1 全体構造と処理フロー")
    add_p("app.pyは以下の7つの主要セクションから構成されており、Streamlitのリプレイスド実行メカニズムに従い、UIの変更検知時に上から順に実行されます。")

    flow_table = doc.add_table(rows=8, cols=3)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(flow_table)
    f_headers = ["コードセクション", "行番号目安", "処理概要・機能"]
    for i, title in enumerate(f_headers):
        set_cell_background(flow_table.rows[0].cells[i], "1F4E79")
        set_cell_margins(flow_table.rows[0].cells[i], top=100, bottom=100, left=120, right=120)
        format_cell_text(flow_table.rows[0].cells[i], title, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    flow_data = [
        ("1. ページ基本設定", "L36 - L82", "画面タイトル、ワイドレイアウト設定、カスタムCSS定義"),
        ("2. データ読み込み処理", "L84 - L120", "文字コード自動判別機能付き関数 `load_csv_data` の定義"),
        ("3. サイドバー制御", "L122 - L173", "作業モード、データソース選択、ファイルアップロード受付"),
        ("4. データ型自動識別", "L175 - L244", "全列、数値列、カテゴリ列、日付列の判定および軸選択コントロール"),
        ("5. メイン画面ヘッダー", "L246 - L280", "データ規模メトリクスカード表示、プレビュー用折りたたみタブ"),
        ("6. メイン機能エリア", "L282 - L823", "作業モードに応じた動的タブ切替（可視化/統計/線形分析/時系列分析）"),
        ("7. フッターガイドライン", "L825 - L835", "各機能の操作説明・活用ガイドの提示")
    ]
    for row_idx, data_tuple in enumerate(flow_data, start=1):
        row_cells = flow_table.rows[row_idx].cells
        bg_color = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data_tuple):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            format_cell_text(row_cells[col_idx], text, bold=(col_idx==0))

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_h2("2.2 データロード機能 (`load_csv_data`) 仕様")
    add_p("日本語環境におけるCSVファイルの読み込み失敗や文字化けを防止するため、試行順序に基づき文字コードを試行するロジックを実装しています。")
    add_bullet("試行エンコーディング順序: ['utf-8', 'shift_jis', 'cp932', 'euc-jp']", "・ 文字コード自動試行: ")
    add_bullet("再試行時に BytesIO の読み込み位置を先頭（0）に復元", "・ ファイルポインタ制御: ")
    add_bullet("@st.cache_data(ttl=3600, show_spinner=False) による計算量・IO削減", "・ キャッシュ制御: ")
    add_bullet("すべての文字コード試行が失敗した場合は None を返し、エラーを防止", "・ エラーハンドリング: ")

    make_callout(
        doc,
        [
            "日本語Windows環境で作成されたShift_JIS/CP932形式のCSVおよびMac/Linux等で作成されたUTF-8形式のCSVの双方に自動対応しています。",
            "ファイルアップロード時およびサンプルデータ読み込み時のいずれでも本共通関数が呼び出されます。"
        ],
        title="📌 文字化け自動判定のポイント"
    )

    add_h2("2.3 データプロファイリング & 列型自動識別")
    add_p("読み込まれたDataFrameの各列は以下のルールに従って自動的に分類されます。")
    add_bullet("DataFrame.select_dtypes(include=[np.number]) に該当する列", "・ 数値列 (numeric_columns): ")
    add_bullet("文字列型 ('object') またはカテゴリ型 ('category') に該当する列", "・ カテゴリ列 (categorical_columns): ")
    add_bullet("非数値列に対して pd.to_datetime() を適用し、60%以上の行が有効な日付時刻に変換できる列", "・ 日付列 (date_columns): ")

    add_h2("2.4 サイドバー制御・操作仕様")
    add_p("サイドバーでは以下の設定を動的に行います。")
    add_bullet("「📈 図を表示する (データ可視化)」「🧪 データを分析する (高度解析)」「📊 両方を表示 (全機能)」の3種類から選択", "1. 作業モードの選択: ")
    add_bullet("「サンプルデータを使用 (sample_data.csv)」または「自分のCSVファイルをアップロード」から選択", "2. データソース選択: ")
    add_bullet("X軸列、Y軸数値列（複数選択可）、グループ分け・色分け列（オプション）を選択", "3. グラフ・分析基本列設定: ")
    add_bullet("図を表示するモード時に、表示したいグラフ（折れ線、棒、散布図、ヒストグラム、箱ひげ図、円グラフ）を個別チェックボックスでON/OFF制御", "4. 表示グラフタイプ選択: ")

    # ---------------------------------------------------------
    # 3. 機能詳細仕様
    # ---------------------------------------------------------
    add_h1("第3章 機能詳細仕様")

    add_h2("3.1 データサマリー & プレビュー表示")
    add_p("メイン画面上部にはデータセット全体の規模を示す4つのメトリクスカードと、折りたたみ（st.expander）形式のプレビューエリアが配置されます。")
    add_bullet("総レコード数（行）、総項目数（列）、数値列数、検出された日付列数", "・ メトリクス表示: ")
    add_bullet("「📄 データプレビュー」「📈 基本記述統計サマリー (describe)」「ℹ️ 列情報・欠損値チェック」の3タブで切り替え表示", "・ プレビュータブ構成: ")

    add_h2("3.2 モジュール1: インタラクティブデータ可視化機能 (📈 図を表示する)")
    add_p("Plotlyを利用した動的グラフを描画します。画面サイズに応じて2列グリッドで配置されます。")

    viz_table = doc.add_table(rows=7, cols=3)
    viz_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(viz_table)
    v_headers = ["グラフ種類", "Plotly関数", "特徴および設定仕様"]
    for i, title in enumerate(v_headers):
        set_cell_background(viz_table.rows[0].cells[i], "1F4E79")
        set_cell_margins(viz_table.rows[0].cells[i], top=100, bottom=100, left=120, right=120)
        format_cell_text(viz_table.rows[0].cells[i], title, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    viz_data = [
        ("📈 折れ線グラフ", "px.line", "markers=True, hovermode='x unified' によりホバー時に同X軸値を一括比較可能"),
        ("📊 棒グラフ", "px.bar", "barmode='group' によるグループ比較対応"),
        ("🟡 散布図", "px.scatter", "2つ目のY軸が選択された場合はマーカーサイズに数値を反映するバブルチャット化対応"),
        ("📶 ヒストグラム", "px.histogram", "marginal='box' を指定し、上部に箱ひげ図を同時描画"),
        ("📦 箱ひげ図", "px.box", "points='all' により全データポイントを表示し外れ値を可視化"),
        ("🥧 円グラフ", "px.pie", "hole=0.3 のドーナツチャート形式。ユニーク数が12以下のカテゴリ列を割り当て")
    ]
    for row_idx, data_tuple in enumerate(viz_data, start=1):
        row_cells = viz_table.rows[row_idx].cells
        bg_color = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data_tuple):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            format_cell_text(row_cells[col_idx], text, bold=(col_idx==0))

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_h2("3.3 モジュール2: 統計学的分析機能 (🧪 データを分析する)")
    add_p("統計学的分析モジュールは以下の3つのサブタブで構成されます。")

    add_h3("サブ機能 2-1: 詳細記述統計量")
    add_p("数値データ列について、通常の基本統計量に加え統計的推定指標を評価・算出します。")
    add_bullet("サンプル数 (N)、平均値、中央値 (50%)、標準偏差 (Std)、分散 (Var)", "・ 基本推計量: ")
    add_bullet("四分位範囲 (IQR = Q3 - Q1)", "・ 頑健性指標: ")
    add_bullet("歪度 (Skewness - 分布の左右の偏り)、尖度 (Kurtosis - 分布の尖り具合)", "・ 形状指標: ")
    add_bullet("t分布に基づく平均値の95%信頼区間 (CI range = stats.t.interval(0.95, df=n-1, loc=mean, scale=sem))", "・ 信頼区間: ")

    add_h3("サブ機能 2-2: 相関分析 (ヒートマップ & 有意性判定)")
    add_p("数値変数間の相関係数を算出し、視覚化と統計的有意性判定を行います。")
    add_bullet("「ピアソン (積率相関 - 直線関係)」または「スピアマン (順位相関 - 単調関係・外れ値耐性)」を選択可能", "・ 相関係数タイプ選択: ")
    add_bullet("px.imshow を使用し、カラーマップ 'RdBu_r' (-1〜1) で可視化", "・ ヒートマップ描画: ")
    add_bullet("scipy.stats.pearsonr または spearmanr により、全変数ペア間のp値を厳密に算出。p < 0.05 を基準に「統計的有意差あり」を一覧判定表示", "・ 有意性 (p値) 評価: ")

    add_h3("サブ機能 2-3: 仮説検定・グループ比較")
    add_p("指定されたカテゴリ列のグループ数に応じて適切な検定手法を適用します。")
    add_bullet("パラメトリック検定「Welchのt検定 (equal_var=False)」およびノンパラメトリック検定「Mann-Whitney U検定」を両方実行・提示", "・ 2群の比較時: ")
    add_bullet("パラメトリック「一元配置分散分析 (ANOVA)」およびノンパラメトリック「Kruskal-Wallis検定」を実行", "・ 3群以上の多群比較時: ")
    add_bullet("自動的にインサイト判定文（「統計的に有意な平均差が見られます」等）を画面出力", "・ インサイト判定: ")

    add_h2("3.4 モジュール3: 線形分析・回帰分析機能 (📐 OLS最小二乗法)")
    add_p("目的変数（Y）と1つ以上の説明変数（X）の関係を最小二乗法（OLS）によりモデル化します。")
    add_bullet("決定係数 (R²)、調整済みR²、RMSE (二乗平均平方根誤差)、モデル全体のP値 (F検定)", "・ モデル評価指標: ")
    add_bullet("各変数および切片 (const) の回帰係数 (B)、標準誤差、t値、p値 (P>|t|)、95%信頼区間、有意性判定", "・ 回帰係数テーブル: ")
    add_bullet("推定されたパラメータより 「Y = β0 + β1*X1 + β2*X2 ...」 の回帰方程式テキストを動的生成", "・ 自動回帰方程式: ")
    add_bullet("単回帰時は散布図上に回帰直線を表示。重回帰時は「実測値 vs 予測値プロット」および「予測値 vs 残差プロット (y=0基準線)」を描画し残差の均一性を検証", "・ 残差・適合度分析グラフ: ")

    add_h2("3.5 モジュール4: 時系列分析機能 (⏳ 時系列分析)")
    add_p("日付・時間軸に沿ったデータの変動パターン分析および将来トレンド予測を行います。")
    add_bullet("スライダーで指定された期間数（ウィンドウサイズ）に基づき、移動平均（SMA）および上下2標準偏差（±2σ）のボリンジャーバンド領域を背景シェーディング付き描画", "・ サブタブ1 (移動平均 & ボリンジャーバンド): ")
    add_bullet("前期比成長率 (% change = pct_change() * 100) の棒グラフおよび累積和 (cumsum) の折れ線グラフを表示", "・ サブタブ2 (変動率・累積和): ")
    add_bullet("経過時間をインデックス（0, 1, 2, ...）化し、LinearRegressionでトレンドモデルを作成。日付間隔の中央値（diff().median()）から将来の日付を自動算出し予測値をプロット・テーブル出力", "・ サブタブ3 (線形トレンド将来予測): ")

    # ---------------------------------------------------------
    # 4. 入出力インターフェース & 例外処理
    # ---------------------------------------------------------
    add_h1("第4章 入出力仕様 & エラーハンドリング")

    add_h2("4.1 入力フォーマット仕様")
    add_bullet("拡張子 .csv のテキストファイル", "・ ファイル形式: ")
    add_bullet("UTF-8, Shift-JIS, CP932, EUC-JP (自動判別)", "・ 文字コード: ")
    add_bullet("1行目をヘッダー（列名）とする表形式データ", "・ データ構造: ")

    add_h2("4.2 同梱サンプルデータ仕様 (`sample_data.csv`)")
    add_p("ユーザーがファイルを用意していない場合に即座に機能を動作確認できるよう、サンプルCSVが用意されています。")
    add_bullet("店舗の売上データ、客数、客単価、地域、日付等を含むダミーデータ", "・ 内容: ")

    add_h2("4.3 エラーハンドリング・ガードロジック")
    add_p("システム全体の安定動作のため、以下の堅牢な例外処理が組み込まれています。")

    err_table = doc.add_table(rows=6, cols=3)
    err_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(err_table)
    e_headers = ["状況・発生条件", "防止メカニズム / 検知コード", "ユーザーへの表示・処理結果"]
    for i, title in enumerate(e_headers):
        set_cell_background(err_table.rows[0].cells[i], "1F4E79")
        set_cell_margins(err_table.rows[0].cells[i], top=100, bottom=100, left=120, right=120)
        format_cell_text(err_table.rows[0].cells[i], title, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    err_data = [
        ("未対応文字コード/破損ファイル", "try-except (UnicodeDecodeError, ParserError)", "次のエンコーディング試行。全滅時は None を返し案内表示"),
        ("ファイル未選択・空データ", "if df is None or df.empty:", "st.info でファイル選択を促し st.stop() で後続処理を停止"),
        ("数値列の不在", "if numeric_columns: の条件分岐", "「数値形式の列が含まれていません」メッセージを表示"),
        ("回帰分析のデータ件数不足", "if len(clean_df) > len(features_x) + 1:", "「分析を実行するための十分なデータ件数がありません」と警告"),
        ("時系列日付データの不足", "if len(ts_df) >= 3:", "「時系列分析を実行するには有効な日付データが3行以上必要」と警告")
    ]
    for row_idx, data_tuple in enumerate(err_data, start=1):
        row_cells = err_table.rows[row_idx].cells
        bg_color = "F9FAFB" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data_tuple):
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=80, bottom=80, left=120, right=120)
            format_cell_text(row_cells[col_idx], text, bold=(col_idx==0))

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Save document
    filename1 = "app_仕様書.docx"
    filename2 = "app_specification.docx"
    doc.save(filename1)
    doc.save(filename2)
    print(f"Successfully generated specification documents: {filename1}, {filename2}")

if __name__ == "__main__":
    create_document()
